import sys
import json
import os
import base64
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def generate_modern_resume(json_path):
    # Load data
    with open(json_path, 'r') as f:
        data = json.load(f)

    document = Document()
    
    # Set narrow margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Create Main Layout Table (1 Row, 2 Columns)
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.5) # Sidebar
    table.columns[1].width = Inches(5.0) # Main Content

    # --- Sidebar (Left Column) ---
    sidebar_cell = table.cell(0, 0)
    
    # Set Sidebar Background Color (Light Teal)
    shading_elm = parse_xml(r'<w:shd {} w:fill="E0F2F1"/>'.format(nsdecls('w')))
    sidebar_cell._tc.get_or_add_tcPr().append(shading_elm)

    # Photo
    if 'photo' in data and data['photo']:
        try:
            header, encoded = data['photo'].split(",", 1)
            image_data = base64.b64decode(encoded)
            image_stream = io.BytesIO(image_data)
            
            p = sidebar_cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(image_stream, width=Inches(1.5))
        except Exception as e:
            print(f"Error processing photo: {e}")

    # Profile Section
    if 'objective' in data and data['objective']:
        p = sidebar_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        run = p.add_run("PROFILE")
        run.bold = True
        run.font.size = Pt(14)
        
        p = sidebar_cell.add_paragraph(data['objective'])
        p.style.font.size = Pt(10)

    # Contact Section
    p = sidebar_cell.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run("CONTACT ME")
    run.bold = True
    run.font.size = Pt(14)

    if 'address' in data and data['address']:
        sidebar_cell.add_paragraph(data['address']).style.font.size = Pt(10)
    if 'phone' in data and data['phone']:
        sidebar_cell.add_paragraph(data['phone']).style.font.size = Pt(10)
    if 'email' in data and data['email']:
        sidebar_cell.add_paragraph(data['email']).style.font.size = Pt(10)

    # Skills Section (Sidebar)
    if 'skills' in data and len(data['skills']) > 0:
        p = sidebar_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        run = p.add_run("MY PRO SKILL")
        run.bold = True
        run.font.size = Pt(14)

        for skill in data['skills']:
            p = sidebar_cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run_name = p.add_run(skill.get('name', ''))
            run_name.bold = True
            run_name.font.size = Pt(10)
            
            # Simple text representation of bar
            level = int(skill.get('level', 50))
            blocks = int(level / 10)
            bar = "█" * blocks + "░" * (10 - blocks)
            p_bar = sidebar_cell.add_paragraph(bar)
            p_bar.style.font.size = Pt(8)

    # --- Main Content (Right Column) ---
    main_cell = table.cell(0, 1)
    
    # Name & Title
    p = main_cell.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    run_name = p.add_run(data.get('name', 'YOUR NAME').upper())
    run_name.bold = True
    run_name.font.size = Pt(28)
    
    p = main_cell.add_paragraph()
    run_title = p.add_run(data.get('jobTitle', 'JOB TITLE').upper())
    run_title.font.size = Pt(12)
    run_title.font.color.rgb = RGBColor(100, 100, 100)
    
    # Education
    if 'education' in data and len(data['education']) > 0:
        p = main_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(30)
        run = p.add_run("EDUCATION")
        run.bold = True
        run.font.size = Pt(16)
        
        for edu in data['education']:
            p = main_cell.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            run_degree = p.add_run(edu.get('degree', '').upper())
            run_degree.bold = True
            run_degree.font.size = Pt(11)
            
            p = main_cell.add_paragraph()
            run_year = p.add_run(edu.get('year', ''))
            run_year.bold = True
            run_year.font.size = Pt(10)
            
            p = main_cell.add_paragraph(edu.get('school', ''))
            p.style.font.size = Pt(10)

    # Experience
    if 'experience' in data and len(data['experience']) > 0:
        p = main_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(30)
        run = p.add_run("EXPERIENCE")
        run.bold = True
        run.font.size = Pt(16)
        
        for exp in data['experience']:
            p = main_cell.add_paragraph()
            p.paragraph_format.space_before = Pt(15)
            run_title = p.add_run(f"{exp.get('title', '').upper()} AT {exp.get('company', '').upper()}")
            run_title.bold = True
            run_title.font.size = Pt(11)
            
            p = main_cell.add_paragraph(exp.get('description', ''))
            p.style.font.size = Pt(10)


    # Save the document
    output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docx')
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    output_filename = os.path.join(output_dir, 'modern_resume.docx')
    document.save(output_filename)
    print(f"Modern Resume generated successfully: {output_filename}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        generate_modern_resume(sys.argv[1])
    else:
        print("Error: No JSON file provided.")
