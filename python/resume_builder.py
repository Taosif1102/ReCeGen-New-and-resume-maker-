import sys
import json
import os
from docx import Document
from docx.shared import Pt

def generate_resume(json_path):
    # Load data
    with open(json_path, 'r') as f:
        data = json.load(f)

    document = Document()

    # Main Header
    document.add_heading(data.get('name', 'Your Name'), 0)

    # Personal Information
    document.add_heading('Personal Information', level=1)
    if 'address' in data: document.add_paragraph(f"Address: {data['address']}")
    if 'phone' in data: document.add_paragraph(f"Phone: {data['phone']}")
    if 'email' in data: document.add_paragraph(f"Email: {data['email']}")
    if 'linkedin' in data: document.add_paragraph(f"LinkedIn: {data['linkedin']}")

    # Career Objective
    if 'objective' in data:
        document.add_heading('Career Objective', level=1)
        document.add_paragraph(data['objective'])

    # Experience
    if 'experience' in data and len(data['experience']) > 0:
        document.add_heading('Experience', level=1)
        for exp in data['experience']:
            p = document.add_paragraph()
            p.add_run(f"{exp.get('title', '')}").bold = True
            p.add_run(f" at {exp.get('company', '')}").italic = True
            p.add_run(f" ({exp.get('date', '')})")
            document.add_paragraph(exp.get('description', ''))

    # Academic Details
    if 'education' in data and len(data['education']) > 0:
        document.add_heading('Academic Details', level=1)
        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Degree'
        hdr_cells[1].text = 'Institute'
        hdr_cells[2].text = 'Year'
        
        for edu in data['education']:
            row_cells = table.add_row().cells
            row_cells[0].text = edu.get('degree', '')
            row_cells[1].text = edu.get('school', '')
            row_cells[2].text = edu.get('year', '')

    # Skills
    if 'skills' in data:
        document.add_heading('Skills', level=1)
        document.add_paragraph(data['skills'])

    # Save the document
    output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docx')
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    output_filename = os.path.join(output_dir, 'resume_output.docx')
    document.save(output_filename)
    print(f"Resume generated successfully: {output_filename}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        generate_resume(sys.argv[1])
    else:
        print("Error: No JSON file provided.")
