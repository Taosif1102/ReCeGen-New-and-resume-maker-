import sys
import json
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_hex):
    """
    Set background color for a table cell.
    """
    cell_properties = cell._element.tcPr
    try:
        cell_shading = cell_properties.xpath('w:shd')[0]
    except IndexError:
        cell_shading = OxmlElement('w:shd')
        cell_properties.append(cell_shading)
    
    cell_shading.set(qn('w:fill'), color_hex)

def create_cv(json_path):
    # Load data
    with open(json_path, 'r') as f:
        data = json.load(f)

    document = Document()

    # Set margins to 0 for full-width sidebar effect (simulated)
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Create a table for the layout (1 row, 2 columns)
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    
    # Set column widths (approx 1/3 and 2/3)
    # Note: Total width is approx 7.5 inches with 0.5 margins
    table.columns[0].width = Inches(2.5) 
    table.columns[1].width = Inches(5.0)

    # --- LEFT COLUMN (SIDEBAR) ---
    left_cell = table.cell(0, 0)
    set_cell_background(left_cell, "1E293B") # Slate-800 hex

    # Helper to add white text to sidebar
    def add_sidebar_text(text, size=10, bold=False, color=RGBColor(255, 255, 255), alignment=WD_ALIGN_PARAGRAPH.LEFT):
        p = left_cell.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.bold = bold
        p.alignment = alignment
        return p

    def add_sidebar_header(text):
        p = left_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text.upper())
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(148, 163, 184) # Slate-400
        run.bold = True
        # Add bottom border simulation (using underscore or just spacing)
        # For now, just spacing
        return p

    # Name & Title
    add_sidebar_text(data.get('name', 'YOUR NAME').upper(), size=20, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_sidebar_text(data.get('title', 'Professional Title'), size=12, color=RGBColor(129, 140, 248), alignment=WD_ALIGN_PARAGRAPH.CENTER) # Indigo-400
    left_cell.add_paragraph() # Spacer

    # Contact
    add_sidebar_header("Contact")
    if 'phone' in data: add_sidebar_text(f"Phone: {data['phone']}")
    if 'email' in data: add_sidebar_text(f"Email: {data['email']}")
    if 'location' in data: add_sidebar_text(f"Loc: {data['location']}")
    
    # Skills
    if 'skills' in data:
        add_sidebar_header("Skills")
        skills_list = data['skills'].split(',') if isinstance(data['skills'], str) else data['skills']
        for skill in skills_list:
            if skill.strip():
                add_sidebar_text(f"• {skill.strip()}")

    # Languages
    if 'languages' in data:
        add_sidebar_header("Languages")
        langs_list = data['languages'].split(',') if isinstance(data['languages'], str) else data['languages']
        for lang in langs_list:
            if lang.strip():
                add_sidebar_text(f"• {lang.strip()}")

    # Hobbies
    if 'hobbies' in data:
        add_sidebar_header("Hobbies")
        hobbies_list = data['hobbies'].split(',') if isinstance(data['hobbies'], str) else data['hobbies']
        for hobby in hobbies_list:
            if hobby.strip():
                add_sidebar_text(f"• {hobby.strip()}")


    # --- RIGHT COLUMN (MAIN CONTENT) ---
    right_cell = table.cell(0, 1)
    
    # Helper for main content
    def add_main_header(text):
        p = right_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text.upper())
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(15, 23, 42) # Slate-900
        run.bold = True
        return p

    # Profile
    if 'summary' in data:
        add_main_header("Profile")
        p = right_cell.add_paragraph(data['summary'])
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Experience
    if 'experience' in data and len(data['experience']) > 0:
        add_main_header("Work Experience")
        for exp in data['experience']:
            # Title & Date
            p = right_cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            title_run = p.add_run(exp.get('title', ''))
            title_run.bold = True
            title_run.font.size = Pt(12)
            
            # Date (Right aligned simulation using tabs or just appending)
            # Simple append for now
            date_run = p.add_run(f"  |  {exp.get('date', '')}")
            date_run.font.size = Pt(10)
            date_run.font.color.rgb = RGBColor(100, 116, 139) # Slate-500

            # Company
            p_comp = right_cell.add_paragraph()
            p_comp.paragraph_format.space_after = Pt(4)
            comp_run = p_comp.add_run(exp.get('company', ''))
            comp_run.italic = True
            comp_run.font.color.rgb = RGBColor(51, 65, 85) # Slate-700

            # Description
            p_desc = right_cell.add_paragraph(exp.get('description', ''))
            p_desc.paragraph_format.space_after = Pt(12)

    # Education
    if 'education' in data and len(data['education']) > 0:
        add_main_header("Education")
        for edu in data['education']:
            p = right_cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            deg_run = p.add_run(edu.get('degree', ''))
            deg_run.bold = True
            deg_run.font.size = Pt(12)

            p_school = right_cell.add_paragraph()
            p_school.paragraph_format.space_after = Pt(12)
            school_run = p_school.add_run(f"{edu.get('school', '')} | {edu.get('year', '')}")
            school_run.italic = True
            school_run.font.color.rgb = RGBColor(51, 65, 85)

    # Save the document
    output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docx')
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    output_path = os.path.join(output_dir, 'cv_output.docx')
    document.save(output_path)
    print(f"CV generated successfully: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        create_cv(sys.argv[1])
    else:
        print("Error: No JSON file provided.")
