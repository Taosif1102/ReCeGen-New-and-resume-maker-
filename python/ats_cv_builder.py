import sys
import json
import os
import base64
import io
from docx import Document
from docx.shared import Pt, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_ats_cv(json_path):
    # Load data
    with open(json_path, 'r') as f:
        data = json.load(f)

    document = Document()
    
    # Set margins (Normal 1 inch is good for ATS, but we can make it slightly smaller if needed)
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Header with Photo (using a table for layout)
    # ATS parsers can struggle with floating images, so a table is safer for layout, 
    # though some very old ATS might ignore the table content. 
    # However, for a "CV with Photo", this is the standard approach.
    
    header_table = document.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Inches(4.5) # Text column
    header_table.columns[1].width = Inches(1.5) # Photo column

    # Text Cell
    text_cell = header_table.cell(0, 0)
    name_paragraph = text_cell.paragraphs[0]
    name_run = name_paragraph.add_run(data.get('name', 'Your Name'))
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.name = 'Arial'

    # Contact Info
    contact_info = []
    if 'address' in data and data['address']: contact_info.append(data['address'])
    if 'phone' in data and data['phone']: contact_info.append(data['phone'])
    if 'email' in data and data['email']: contact_info.append(data['email'])
    if 'linkedin' in data and data['linkedin']: contact_info.append(data['linkedin'])
    
    if contact_info:
        contact_p = text_cell.add_paragraph(" | ".join(contact_info))
        contact_p.style.font.name = 'Arial'
        contact_p.style.font.size = Pt(10)

    # Photo Cell
    photo_cell = header_table.cell(0, 1)
    if 'photo' in data and data['photo']:
        try:
            # Decode base64 image
            # Expecting data:image/png;base64,......
            header, encoded = data['photo'].split(",", 1)
            image_data = base64.b64decode(encoded)
            image_stream = io.BytesIO(image_data)
            
            paragraph = photo_cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run()
            run.add_picture(image_stream, width=Inches(1.2))
        except Exception as e:
            print(f"Error processing photo: {e}")

    document.add_paragraph() # Spacer

    # Professional Summary
    if 'objective' in data and data['objective']:
        h = document.add_heading('PROFESSIONAL SUMMARY', level=1)
        h.style.font.name = 'Arial'
        h.style.font.color.rgb = None # Default black
        p = document.add_paragraph(data['objective'])
        p.style.font.name = 'Arial'

    # Experience
    if 'experience' in data and len(data['experience']) > 0:
        h = document.add_heading('WORK EXPERIENCE', level=1)
        h.style.font.name = 'Arial'
        h.style.font.color.rgb = None
        
        for exp in data['experience']:
            p = document.add_paragraph()
            p.style.font.name = 'Arial'
            
            # Title and Company
            title_run = p.add_run(f"{exp.get('title', '')}")
            title_run.bold = True
            title_run.font.size = Pt(12)
            
            company_run = p.add_run(f" | {exp.get('company', '')}")
            company_run.bold = True
            
            # Date
            if 'date' in exp:
                p.add_run(f"\t{exp.get('date', '')}").bold = False
            
            # Description
            if 'description' in exp:
                desc_p = document.add_paragraph(exp.get('description', ''))
                desc_p.style.font.name = 'Arial'
                desc_p.paragraph_format.left_indent = Pt(10)

    # Skills
    if 'skills' in data and data['skills']:
        h = document.add_heading('SKILLS', level=1)
        h.style.font.name = 'Arial'
        h.style.font.color.rgb = None
        p = document.add_paragraph(data['skills'])
        p.style.font.name = 'Arial'

    # Education
    if 'education' in data and len(data['education']) > 0:
        h = document.add_heading('EDUCATION', level=1)
        h.style.font.name = 'Arial'
        h.style.font.color.rgb = None
        
        for edu in data['education']:
            p = document.add_paragraph()
            p.style.font.name = 'Arial'
            
            school_run = p.add_run(f"{edu.get('school', '')}")
            school_run.bold = True
            
            if 'degree' in edu:
                p.add_run(f" - {edu.get('degree', '')}")
                
            if 'year' in edu:
                p.add_run(f"\t{edu.get('year', '')}")

    # Save the document
    output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docx')
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    output_filename = os.path.join(output_dir, 'ats_cv_output.docx')
    document.save(output_filename)
    print(f"ATS CV generated successfully: {output_filename}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        generate_ats_cv(sys.argv[1])
    else:
        print("Error: No JSON file provided.")
