import sys
import json
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_ats_resume(json_path):
    # Load data
    with open(json_path, 'r') as f:
        data = json.load(f)

    document = Document()
    
    # Set margins (Standard 1 inch)
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style Configuration
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Header (Name & Contact)
    name_p = document.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(data.get('name', 'Your Name'))
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.name = 'Arial'

    # Contact Info (Single line, separated by |)
    contact_info = []
    if 'address' in data and data['address']: contact_info.append(data['address'])
    if 'phone' in data and data['phone']: contact_info.append(data['phone'])
    if 'email' in data and data['email']: contact_info.append(data['email'])
    if 'linkedin' in data and data['linkedin']: contact_info.append(data['linkedin'])
    
    if contact_info:
        contact_p = document.add_paragraph(" | ".join(contact_info))
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.style.font.name = 'Arial'

    document.add_paragraph() # Spacer

    # Helper for Section Headings
    def add_section_heading(text):
        h = document.add_heading(text.upper(), level=1)
        h.style.font.name = 'Arial'
        h.style.font.size = Pt(14)
        h.style.font.color.rgb = None # Black
        h.style.font.bold = True
        # Add border bottom manually if needed, but standard heading style is usually fine for ATS

    # Professional Summary
    if 'objective' in data and data['objective']:
        add_section_heading('Professional Summary')
        p = document.add_paragraph(data['objective'])
        p.style.font.name = 'Arial'

    # Work Experience
    if 'experience' in data and len(data['experience']) > 0:
        add_section_heading('Work Experience')
        
        for exp in data['experience']:
            # Title
            p_title = document.add_paragraph()
            p_title.paragraph_format.space_after = Pt(2)
            run_title = p_title.add_run(f"{exp.get('title', '')}")
            run_title.bold = True
            run_title.font.size = Pt(12)
            
            # Company & Date line
            p_meta = document.add_paragraph()
            p_meta.paragraph_format.space_before = Pt(0)
            p_meta.paragraph_format.space_after = Pt(2)
            run_company = p_meta.add_run(f"{exp.get('company', '')}")
            run_company.bold = True
            
            if 'date' in exp:
                p_meta.add_run(f" | {exp.get('date', '')}")

            # Description
            if 'description' in exp:
                desc_p = document.add_paragraph(exp.get('description', ''))
                desc_p.style.font.name = 'Arial'
                desc_p.paragraph_format.space_after = Pt(10)

    # Education (No Tables!)
    if 'education' in data and len(data['education']) > 0:
        add_section_heading('Education')
        
        for edu in data['education']:
            p = document.add_paragraph()
            
            # School Name
            run_school = p.add_run(f"{edu.get('school', '')}")
            run_school.bold = True
            
            # Degree and Year
            details = []
            if 'degree' in edu: details.append(edu.get('degree', ''))
            if 'year' in edu: details.append(edu.get('year', ''))
            
            if details:
                p.add_run(f" | {' - '.join(details)}")

    # Skills
    if 'skills' in data and data['skills']:
        add_section_heading('Skills')
        p = document.add_paragraph(data['skills'])
        p.style.font.name = 'Arial'

    # Save the document
    output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docx')
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    output_filename = os.path.join(output_dir, 'ats_resume_output.docx')
    document.save(output_filename)
    print(f"ATS Resume generated successfully: {output_filename}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        generate_ats_resume(sys.argv[1])
    else:
        print("Error: No JSON file provided.")
